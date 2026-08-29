# -*- coding: utf-8 -*-
"""
GNSS_COMMENT12_MANUSCRIPT_NUMERICAL_AUDIT_v1_1.py
==================================================

Reviewer Comment 12 — Manuscript-wide numerical consistency audit
v1.1

Purpose
-------
Read the current manuscript DOCX in read-only mode and compare numerical
correlation claims against the final evidence registry.

v1.1 fixes
----------
1) Only fragments containing an actual rho/p/q statistic can become a
   numerical discrepancy. Metric-only figure/caption references are no longer
   marked REPLACE.

2) Multi-metric sentences are parsed using metric-local windows so statements
   such as:
       CMC P95 ... rho=...
       CMC robust ... rho=...
       C/N0 ... rho=...
   can be resolved separately.

3) Abstract and Conclusions are handled reliably even when multiple
   association claims are reported in one paragraph.

4) 2D/horizontal RMSE and 3D RMSE are explicitly separated.

5) Legacy q-based primary claims are classified as:
       LEGACY_FDR_VALUE_VERIFY_OR_REPLACE
   rather than silently treated as equivalent to the new exact-p primary lock.

6) The output ledger is organized around publication actions:
       KEEP
       REPLACE
       RELABEL
       VERIFY_SOURCE
       NO_ACTION

Primary source-of-truth
-----------------------
Product-adjusted + exact Test-blocked permutation.

Locked primary values:
    CMC_P95    -> horizontal / 2D RMSE
    CMC_ROBUST -> horizontal / 2D RMSE
    CN0        -> horizontal / 2D RMSE

q-value policy
--------------
The Comment-11 primary source defines exact blocked-permutation p-values,
not FDR-adjusted q-values. A legacy q-value attached to a primary claim must
either:
    a) be explicitly linked to its original finalized FDR source and clearly
       relabelled as secondary/legacy evidence, or
    b) be replaced in the primary narrative by the locked exact-p result.

Other claims
------------
Product sensitivity, 3D-RMSE correlations, product-fixed-effect partial R^2,
rank correlations, satellite-count correlations, etc. are not overwritten by
the Comment-12 primary lock. They are marked VERIFY_SOURCE unless already
linked to another authoritative source.

Expected manuscript
-------------------
Sensors_Ready_Master_Manuscript_v2_1_SUBMISSION_CANDIDATE_r1.docx

Expected registry
-----------------
C:\IEEE\GNSS_ANALYSIS\COMMENT12_NUMERICAL_CONSISTENCY_AUDIT_V1\
    COMMENT12_FINAL_EVIDENCE_REGISTRY.csv

Outputs
-------
C:\IEEE\GNSS_ANALYSIS\COMMENT12_MANUSCRIPT_NUMERICAL_AUDIT_V1_1\

    COMMENT12_V1_1_STATISTIC_INVENTORY.csv
    COMMENT12_V1_1_CLAIM_LEDGER.csv
    COMMENT12_V1_1_PRIMARY_CLAIM_OCCURRENCES.csv
    COMMENT12_V1_1_SECTION_CONSISTENCY_MATRIX.csv
    COMMENT12_V1_1_REVISION_ACTIONS.csv
    COMMENT12_V1_1_MANUSCRIPT_NUMERICAL_AUDIT_SUMMARY.txt
    COMMENT12_V1_1_BLOCKERS.csv              # only if needed

The manuscript is never modified.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional, List, Dict
import csv
import math
import re

import numpy as np
import pandas as pd
from docx import Document


# ======================================================================
# CONFIG
# ======================================================================

ROOT = Path(r"C:\IEEE")
MANUSCRIPT_NAME = "Sensors_Ready_Master_Manuscript_v2_1_SUBMISSION_CANDIDATE_r1.docx"

MANUSCRIPT_CANDIDATES = [
    ROOT / MANUSCRIPT_NAME,
    ROOT / "YAZIM" / MANUSCRIPT_NAME,
    Path.cwd() / MANUSCRIPT_NAME,
    Path(__file__).resolve().parent / MANUSCRIPT_NAME,
    Path("/mnt/data") / MANUSCRIPT_NAME,
]

REGISTRY = (
    ROOT / "GNSS_ANALYSIS"
    / "COMMENT12_NUMERICAL_CONSISTENCY_AUDIT_V1"
    / "COMMENT12_FINAL_EVIDENCE_REGISTRY.csv"
)

OUT = (
    ROOT / "GNSS_ANALYSIS"
    / "COMMENT12_MANUSCRIPT_NUMERICAL_AUDIT_V1_1"
)
OUT.mkdir(parents=True, exist_ok=True)

EXPECTED_PRIMARY_METRICS = ["CMC_P95", "CMC_ROBUST", "CN0"]

RHO_TOL = 0.0015
P_TOL = 0.00055

PRIMARY_NARRATIVE_SECTIONS = {
    "Abstract", "Results", "Discussion", "Conclusions"
}

# Metric-local window length after metric mention.
LOCAL_WINDOW_CHARS = 260

METRIC_PATTERNS = {
    "CMC_P95": [
        r"GPS\s*L1\s*CMC\s*P95",
        r"CMC\s*P95",
    ],
    "CMC_ROBUST": [
        r"GPS\s*L1\s*CMC\s*robust\s*(?:dispersion|sigma|σ)",
        r"CMC\s*robust\s*(?:dispersion|sigma|σ)",
        r"robust\s*CMC\s*(?:dispersion|sigma|σ)",
    ],
    "CN0": [
        r"median\s*C\s*/?\s*N[₀0]",
        r"C\s*/?\s*N[₀0]",
        r"CN0",
        r"CNO",
    ],
}

OUTCOME_PATTERNS = {
    "RMSE2D": [
        r"median\s*horizontal\s*RMSE",
        r"horizontal\s*RMSE",
        r"2D\s*RMSE",
        r"2-D\s*RMSE",
        r"horizontal\s*error",
    ],
    "RMSE3D": [
        r"median\s*3D\s*RMSE",
        r"3D\s*RMSE",
        r"three[-\s]*dimensional\s*RMSE",
    ],
    "PRODUCT_SENSITIVITY": [
        r"product\s*sensitivity",
        r"product\s*range",
        r"across[-\s]*product\s*(?:horizontal\s*)?RMSE\s*range",
        r"across[-\s]*product\s*range",
    ],
    "SATELLITE_COUNT": [
        r"usable[-\s]*satellite\s*count",
        r"satellite\s*count",
    ],
    "RANK": [
        r"rank\s*(?:association|correlation)",
        r"Kendall\s*τ",
    ],
}

RHO_RE = re.compile(
    r"(?:Spearman\s*)?(?:ρ|rho)\s*=\s*([+-]?(?:\d+(?:\.\d*)?|\.\d+))",
    re.I
)
P_RE = re.compile(
    r"(?<![A-Za-z])p\s*=\s*([0-9]*\.?[0-9]+(?:[eE][+-]?\d+)?)",
    re.I
)
Q_RE = re.compile(
    r"(?<![A-Za-z])q\s*=\s*([0-9]*\.?[0-9]+(?:[eE][+-]?\d+)?)",
    re.I
)
R2_RE = re.compile(
    r"(?:partial\s*)?R(?:²|\^?2)\s*=\s*([0-9]*\.?[0-9]+)",
    re.I
)


# ======================================================================
# HELPERS
# ======================================================================

def write_csv(path: Path, rows: List[dict], fields: Optional[List[str]] = None):
    if fields is None:
        fields = list(rows[0].keys()) if rows else ["Status"]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        w = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)


def finite(v):
    try:
        return math.isfinite(float(v))
    except Exception:
        return False


def ffloat(v):
    try:
        return float(v)
    except Exception:
        return float("nan")


def approx(a, b, tol):
    return finite(a) and finite(b) and abs(float(a) - float(b)) <= tol


def clean_text(s):
    s = str(s or "")
    replacements = {
        "\u00ad": "",
        "\u2011": "-",
        "\u2013": "-",
        "\u2014": "-",
        "\ufeff": "",
        "\u2080": "0",
    }
    for a, b in replacements.items():
        s = s.replace(a, b)
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def locate_manuscript():
    for p in MANUSCRIPT_CANDIDATES:
        try:
            if p.exists() and p.stat().st_size > 0:
                return p
        except Exception:
            pass
    raise FileNotFoundError(
        "Manuscript not found. Tried:\n" +
        "\n".join(str(p) for p in MANUSCRIPT_CANDIDATES)
    )


# ======================================================================
# REGISTRY
# ======================================================================

def load_registry():
    if not REGISTRY.exists():
        raise FileNotFoundError(REGISTRY)

    reg = pd.read_csv(REGISTRY)

    primary = reg[
        reg["Evidence_Role"].astype(str).eq("PRIMARY_INFERENCE")
    ].copy()

    robustness = reg[
        reg["Evidence_Role"].astype(str).eq("ROBUSTNESS_EVIDENCE")
        & reg["Statistical_Test"].astype(str).str.contains(
            "exact Test-blocked permutation", case=False, na=False
        )
    ].copy()

    if set(primary["Metric"].astype(str)) != set(EXPECTED_PRIMARY_METRICS):
        raise RuntimeError(
            "Unexpected primary metric set: " +
            str(sorted(primary["Metric"].astype(str).tolist()))
        )

    pmap = {}
    for _, r in primary.iterrows():
        pmap[str(r["Metric"])] = {
            "rho": ffloat(r["Rho"]),
            "p": ffloat(r["P"]),
            "claim_id": str(r["Claim_ID"]),
            "source_csv": str(r["Source_CSV"]),
            "source_row": str(r["Source_Row"]),
        }

    rmap = {}
    for _, r in robustness.iterrows():
        m = str(r["Metric"])
        if m in EXPECTED_PRIMARY_METRICS:
            rmap[m] = {
                "rho": ffloat(r["Rho"]),
                "p": ffloat(r["P"]),
                "claim_id": str(r["Claim_ID"]),
            }

    return reg, pmap, rmap


# ======================================================================
# DOCX EXTRACTION
# ======================================================================

def is_heading(p):
    try:
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            return True
    except Exception:
        pass

    t = clean_text(p.text)
    if t in {
        "Abstract", "Highlights", "References", "Supplementary Materials",
        "Author Contributions", "Funding", "Institutional Review Board Statement",
        "Informed Consent Statement", "Data Availability Statement",
        "Acknowledgments", "Conflicts of Interest"
    }:
        return True

    return bool(re.match(r"^\d+(?:\.\d+)*\.\s+\S", t))


def coarse_section(heading):
    h = clean_text(heading)
    hl = h.lower()

    if hl == "abstract":
        return "Abstract"
    if h.startswith("1.") or hl.startswith("introduction"):
        return "Introduction"
    if h.startswith("2.") or "materials and methods" in hl:
        return "Methods"
    if h.startswith("3.") or hl.startswith("results"):
        return "Results"
    if h.startswith("4.") or hl.startswith("discussion"):
        return "Discussion"
    if h.startswith("5.") or hl.startswith("conclusions"):
        return "Conclusions"
    if hl.startswith("supplementary"):
        return "Supplementary"
    if hl.startswith("references"):
        return "References"
    return ""


def extract_units(path):
    doc = Document(path)

    units = []
    current_heading = ""
    current_section = ""

    for i, p in enumerate(doc.paragraphs, start=1):
        txt = clean_text(p.text)
        if not txt:
            continue

        if is_heading(p):
            current_heading = txt
            cs = coarse_section(txt)
            if cs:
                current_section = cs

        units.append({
            "Location_Type": "Paragraph",
            "Location_ID": f"P{i}",
            "Paragraph_Index": i,
            "Table_Index": "",
            "Cell": "",
            "Section": current_section or "Unknown",
            "Subsection": current_heading,
            "Text": txt,
        })

    for ti, table in enumerate(doc.tables, start=1):
        for ri, row in enumerate(table.rows, start=1):
            for ci, cell in enumerate(row.cells, start=1):
                txt = clean_text(cell.text)
                if not txt:
                    continue
                units.append({
                    "Location_Type": "TableCell",
                    "Location_ID": f"T{ti}R{ri}C{ci}",
                    "Paragraph_Index": "",
                    "Table_Index": ti,
                    "Cell": f"R{ri}C{ci}",
                    "Section": "Table",
                    "Subsection": f"Table {ti}",
                    "Text": txt,
                })

    return units


# ======================================================================
# STAT EXTRACTION
# ======================================================================

def extract_stats(text):
    return {
        "rhos": [(m.start(), ffloat(m.group(1))) for m in RHO_RE.finditer(text)],
        "ps": [(m.start(), ffloat(m.group(1))) for m in P_RE.finditer(text)],
        "qs": [(m.start(), ffloat(m.group(1))) for m in Q_RE.finditer(text)],
        "r2s": [(m.start(), ffloat(m.group(1))) for m in R2_RE.finditer(text)],
    }


def metric_mentions(text):
    out = []
    for metric, pats in METRIC_PATTERNS.items():
        for pat in pats:
            for m in re.finditer(pat, text, re.I):
                out.append({
                    "metric": metric,
                    "start": m.start(),
                    "end": m.end(),
                    "match": m.group(0),
                })
    return sorted(out, key=lambda x: x["start"])


def detect_outcome(text):
    found = []
    for out, pats in OUTCOME_PATTERNS.items():
        for pat in pats:
            if re.search(pat, text, re.I):
                found.append(out)
                break

    if "RMSE3D" in found and "RMSE2D" not in found:
        return "RMSE3D"
    if "PRODUCT_SENSITIVITY" in found:
        return "PRODUCT_SENSITIVITY"
    if "SATELLITE_COUNT" in found:
        return "SATELLITE_COUNT"
    if "RANK" in found:
        return "RANK"
    if "RMSE2D" in found:
        return "RMSE2D"
    if "RMSE3D" in found:
        return "RMSE3D"
    return "UNRESOLVED"


def nearest_stat_after(stats, metric_end, next_metric_start, max_chars=LOCAL_WINDOW_CHARS):
    """
    Find rho/p/q in the local region belonging to one metric mention.
    """
    upper = min(
        metric_end + max_chars,
        next_metric_start if next_metric_start is not None else 10**9
    )

    def pick(values):
        candidates = [(pos, val) for pos, val in values if metric_end <= pos < upper]
        if not candidates:
            return float("nan")
        candidates.sort(key=lambda x: x[0])
        return candidates[0][1]

    return {
        "rho": pick(stats["rhos"]),
        "p": pick(stats["ps"]),
        "q": pick(stats["qs"]),
    }


def local_claims_from_unit(text):
    """
    Returns only claims with an actual local rho/p/q.
    This eliminates the v1 false positives from metric-only prose.
    """
    mentions = metric_mentions(text)
    stats = extract_stats(text)

    claims = []

    for i, mm in enumerate(mentions):
        next_start = mentions[i+1]["start"] if i+1 < len(mentions) else None

        local = nearest_stat_after(
            stats,
            mm["end"],
            next_start,
            max_chars=LOCAL_WINDOW_CHARS
        )

        if not any(finite(local[k]) for k in ("rho", "p", "q")):
            continue

        upper = min(
            mm["end"] + LOCAL_WINDOW_CHARS,
            next_start if next_start is not None else len(text)
        )
        window = text[mm["start"]:upper]

        outcome = detect_outcome(window)

        claims.append({
            "Metric": mm["metric"],
            "Metric_Text": mm["match"],
            "Metric_Start": mm["start"],
            "Window": window,
            "Outcome": outcome,
            "Found_Rho": local["rho"],
            "Found_p": local["p"],
            "Found_q": local["q"],
        })

    return claims


# ======================================================================
# DECISION ENGINE
# ======================================================================

def decide(metric, outcome, rho, p, q, section, pmap, rmap):
    exp = pmap[metric]
    rob = rmap.get(metric)

    rho_primary = approx(rho, exp["rho"], RHO_TOL)
    p_primary = approx(p, exp["p"], P_TOL)

    rho_rob = rob is not None and approx(rho, rob["rho"], RHO_TOL)
    p_rob = rob is not None and approx(p, rob["p"], P_TOL)

    # Non-primary outcomes are not overwritten.
    if outcome == "RMSE3D":
        return {
            "Decision": "VERIFY_SOURCE_OTHER_CLAIM",
            "Publication_Action": "VERIFY_SOURCE",
            "Problem_Type": "3D-RMSE association is outside the current horizontal-RMSE primary lock.",
            "Recommended_Action": (
                "Verify against the final 3D-RMSE evidence source. "
                "Do not substitute the horizontal-RMSE lock."
            ),
        }

    if outcome in {"PRODUCT_SENSITIVITY", "SATELLITE_COUNT", "RANK"}:
        return {
            "Decision": "VERIFY_SOURCE_OTHER_CLAIM",
            "Publication_Action": "VERIFY_SOURCE",
            "Problem_Type": f"{outcome} claim is outside the current primary lock.",
            "Recommended_Action": "Verify against its own final evidence source.",
        }

    if outcome == "UNRESOLVED":
        return {
            "Decision": "VERIFY_SOURCE_OTHER_CLAIM",
            "Publication_Action": "VERIFY_SOURCE",
            "Problem_Type": "Outcome could not be resolved from the local metric window.",
            "Recommended_Action": "Review surrounding prose before revision.",
        }

    # From here: primary horizontal/2D claim.
    if finite(q):
        if rho_rob:
            return {
                "Decision": "LEGACY_FDR_VALUE_VERIFY_OR_REPLACE",
                "Publication_Action": "REPLACE",
                "Problem_Type": (
                    "Legacy unadjusted/FDR rho-q pair is being used in the primary horizontal-RMSE narrative."
                ),
                "Recommended_Action": (
                    f"Primary text: rho={exp['rho']:.3f}, exact p={exp['p']:.4f}. "
                    "Retain the legacy unadjusted q-result only if explicitly relabelled as "
                    "secondary robustness evidence and linked to its finalized FDR source."
                ),
            }

        if rho_primary:
            return {
                "Decision": "LEGACY_FDR_VALUE_VERIFY_OR_REPLACE",
                "Publication_Action": "RELABEL",
                "Problem_Type": (
                    "Primary rho matches current lock, but significance is reported as legacy q rather than exact p."
                ),
                "Recommended_Action": (
                    f"Use exact p={exp['p']:.4f} in the primary narrative. "
                    "Any q-value requires an explicitly identified FDR source."
                ),
            }

        return {
            "Decision": "LEGACY_FDR_VALUE_VERIFY_OR_REPLACE",
            "Publication_Action": "REPLACE",
            "Problem_Type": "Primary claim uses an old rho/q pair inconsistent with current primary lock.",
            "Recommended_Action": (
                f"Replace primary narrative with rho={exp['rho']:.3f}, exact p={exp['p']:.4f}. "
                "Verify the old q-result separately before retaining it anywhere."
            ),
        }

    if rho_primary and p_primary:
        return {
            "Decision": "MATCH_PRIMARY_LOCK",
            "Publication_Action": "KEEP",
            "Problem_Type": "",
            "Recommended_Action": "Keep.",
        }

    if rho_rob and (not finite(p) or p_rob):
        return {
            "Decision": "ROBUSTNESS_VALUE_USED_AS_PRIMARY",
            "Publication_Action": "REPLACE",
            "Problem_Type": "Unadjusted robustness result used as primary inference.",
            "Recommended_Action": (
                f"Replace with product-adjusted rho={exp['rho']:.3f}, "
                f"exact p={exp['p']:.4f}."
            ),
        }

    if rho_primary and not finite(p):
        return {
            "Decision": "PRIMARY_P_MISSING",
            "Publication_Action": "RELABEL",
            "Problem_Type": "Primary rho matches, but exact blocked-permutation p is absent.",
            "Recommended_Action": (
                f"Add exact p={exp['p']:.4f}."
            ),
        }

    return {
        "Decision": "REPLACE_WITH_PRIMARY_LOCK",
        "Publication_Action": "REPLACE",
        "Problem_Type": "Primary horizontal-RMSE claim differs from authoritative lock.",
        "Recommended_Action": (
            f"Use rho={exp['rho']:.3f}, exact p={exp['p']:.4f}."
        ),
    }


# ======================================================================
# OTHER STATISTICAL CLAIMS
# ======================================================================

def inventory_other_stats(unit):
    """
    Inventory statistical units not captured as local locked-metric claims.
    """
    text = unit["Text"]
    stats = extract_stats(text)
    out = []

    has_any = any(stats[k] for k in stats)
    if not has_any:
        return out

    locked_claims = local_claims_from_unit(text)
    locked_positions = {c["Metric_Start"] for c in locked_claims}

    # If no locked metrics with stats, classify whole statistical unit as other.
    if not locked_claims:
        out.append({
            **unit,
            "Metric": "",
            "Outcome": detect_outcome(text),
            "Found_Rho": "|".join(str(v) for _, v in stats["rhos"]),
            "Found_p": "|".join(str(v) for _, v in stats["ps"]),
            "Found_q": "|".join(str(v) for _, v in stats["qs"]),
            "Found_R2": "|".join(str(v) for _, v in stats["r2s"]),
            "Decision": "VERIFY_SOURCE_OTHER_CLAIM",
            "Publication_Action": "VERIFY_SOURCE",
            "Problem_Type": "Statistical claim outside Comment-12 primary lock.",
            "Recommended_Action": "Verify against its own final evidence source.",
        })

    return out


# ======================================================================
# MAIN
# ======================================================================

def main():
    blockers = []

    try:
        manuscript = locate_manuscript()
        reg, pmap, rmap = load_registry()
    except Exception as e:
        blockers.append({"Issue": "INITIALIZATION_FAIL", "Detail": str(e)})
        write_csv(
            OUT / "COMMENT12_V1_1_BLOCKERS.csv",
            blockers,
            ["Issue", "Detail"]
        )
        raise

    units = extract_units(manuscript)

    statistic_inventory = []
    claim_rows = []
    other_rows = []

    audit_id = 1

    for unit in units:
        text = unit["Text"]
        stats = extract_stats(text)

        if not any(stats[k] for k in stats):
            continue

        statistic_inventory.append({
            **unit,
            "Rho_Values": "|".join(str(v) for _, v in stats["rhos"]),
            "P_Values": "|".join(str(v) for _, v in stats["ps"]),
            "Q_Values": "|".join(str(v) for _, v in stats["qs"]),
            "R2_Values": "|".join(str(v) for _, v in stats["r2s"]),
        })

        local_claims = local_claims_from_unit(text)

        if local_claims:
            for c in local_claims:
                dec = decide(
                    c["Metric"],
                    c["Outcome"],
                    c["Found_Rho"],
                    c["Found_p"],
                    c["Found_q"],
                    unit["Section"],
                    pmap,
                    rmap,
                )

                exp = pmap[c["Metric"]]

                claim_rows.append({
                    "Audit_ID": f"C12M11-{audit_id:04d}",
                    **unit,
                    "Metric": c["Metric"],
                    "Metric_Text": c["Metric_Text"],
                    "Outcome": c["Outcome"],
                    "Local_Window": c["Window"],
                    "Found_Rho": c["Found_Rho"] if finite(c["Found_Rho"]) else "",
                    "Found_p": c["Found_p"] if finite(c["Found_p"]) else "",
                    "Found_q": c["Found_q"] if finite(c["Found_q"]) else "",
                    "Expected_Rho": exp["rho"],
                    "Expected_p": exp["p"],
                    "Expected_Claim_ID": exp["claim_id"],
                    "Expected_Source_CSV": exp["source_csv"],
                    "Expected_Source_Row": exp["source_row"],
                    **dec,
                })
                audit_id += 1

        # Other stats not associated with locked local metric claims
        if not local_claims:
            for r in inventory_other_stats(unit):
                r["Audit_ID"] = f"C12M11-{audit_id:04d}"
                audit_id += 1
                other_rows.append(r)

    all_rows = claim_rows + other_rows

    # ------------------------------------------------------------------
    # Primary occurrences
    # ------------------------------------------------------------------
    primary_occ = [
        r for r in claim_rows
        if r["Outcome"] == "RMSE2D"
        and r["Section"] in PRIMARY_NARRATIVE_SECTIONS
    ]

    # ------------------------------------------------------------------
    # Section consistency matrix
    # ------------------------------------------------------------------
    matrix = []

    for metric in EXPECTED_PRIMARY_METRICS:
        exp = pmap[metric]

        for section in ["Abstract", "Results", "Discussion", "Conclusions"]:
            g = [
                r for r in primary_occ
                if r["Metric"] == metric and r["Section"] == section
            ]

            if not g:
                matrix.append({
                    "Metric": metric,
                    "Section": section,
                    "Occurrence_N": 0,
                    "Expected_Rho": exp["rho"],
                    "Expected_p": exp["p"],
                    "Found_Rho": "",
                    "Found_p": "",
                    "Found_q": "",
                    "Actions": "",
                    "Status": "MISSING_OR_NOT_NUMERICALLY_STATED",
                })
                continue

            actions = sorted(set(r["Publication_Action"] for r in g))
            statuses = sorted(set(r["Decision"] for r in g))

            found_rho = sorted(set(str(r["Found_Rho"]) for r in g if r["Found_Rho"] != ""))
            found_p = sorted(set(str(r["Found_p"]) for r in g if r["Found_p"] != ""))
            found_q = sorted(set(str(r["Found_q"]) for r in g if r["Found_q"] != ""))

            pass_section = all(r["Publication_Action"] == "KEEP" for r in g)

            matrix.append({
                "Metric": metric,
                "Section": section,
                "Occurrence_N": len(g),
                "Expected_Rho": exp["rho"],
                "Expected_p": exp["p"],
                "Found_Rho": "|".join(found_rho),
                "Found_p": "|".join(found_p),
                "Found_q": "|".join(found_q),
                "Actions": "|".join(actions),
                "Decisions": "|".join(statuses),
                "Status": "PASS_PRIMARY_LOCK" if pass_section else "REVISE_SECTION",
            })

    # ------------------------------------------------------------------
    # Revision action table
    # ------------------------------------------------------------------
    action_rows = [
        r for r in claim_rows
        if r["Publication_Action"] in {"REPLACE", "RELABEL", "VERIFY_SOURCE"}
    ]

    # ------------------------------------------------------------------
    # Save outputs
    # ------------------------------------------------------------------
    pd.DataFrame(statistic_inventory).to_csv(
        OUT / "COMMENT12_V1_1_STATISTIC_INVENTORY.csv",
        index=False,
        encoding="utf-8-sig"
    )

    pd.DataFrame(all_rows).to_csv(
        OUT / "COMMENT12_V1_1_CLAIM_LEDGER.csv",
        index=False,
        encoding="utf-8-sig"
    )

    pd.DataFrame(primary_occ).to_csv(
        OUT / "COMMENT12_V1_1_PRIMARY_CLAIM_OCCURRENCES.csv",
        index=False,
        encoding="utf-8-sig"
    )

    pd.DataFrame(matrix).to_csv(
        OUT / "COMMENT12_V1_1_SECTION_CONSISTENCY_MATRIX.csv",
        index=False,
        encoding="utf-8-sig"
    )

    pd.DataFrame(action_rows).to_csv(
        OUT / "COMMENT12_V1_1_REVISION_ACTIONS.csv",
        index=False,
        encoding="utf-8-sig"
    )

    # ------------------------------------------------------------------
    # Summary
    # ------------------------------------------------------------------
    cdf = pd.DataFrame(claim_rows)
    odf = pd.DataFrame(other_rows)

    summary = OUT / "COMMENT12_V1_1_MANUSCRIPT_NUMERICAL_AUDIT_SUMMARY.txt"

    with summary.open("w", encoding="utf-8") as f:
        f.write("GNSS Comment 12 — Manuscript Numerical Audit v1.1\n")
        f.write("=" * 104 + "\n\n")

        f.write(f"Manuscript : {manuscript}\n")
        f.write(f"Registry   : {REGISTRY}\n")
        f.write(f"Units read : {len(units)}\n")
        f.write(f"Statistic-containing units : {len(statistic_inventory)}\n")
        f.write(f"Locked-metric numerical claims : {len(claim_rows)}\n")
        f.write(f"Other statistical claim units  : {len(other_rows)}\n\n")

        f.write("AUTHORITATIVE PRIMARY LOCK\n")
        f.write("-" * 104 + "\n")
        for metric in EXPECTED_PRIMARY_METRICS:
            e = pmap[metric]
            f.write(
                f"{metric:10s} -> horizontal RMSE | "
                f"rho={e['rho']:.6f} | exact p={e['p']:.6f}\n"
            )

        if not cdf.empty:
            f.write("\nLOCKED-METRIC DECISION COUNTS\n")
            f.write("-" * 104 + "\n")
            for dec, n in cdf["Decision"].value_counts().items():
                f.write(f"{dec}: {n}\n")

            f.write("\nPUBLICATION ACTION COUNTS\n")
            f.write("-" * 104 + "\n")
            for action, n in cdf["Publication_Action"].value_counts().items():
                f.write(f"{action}: {n}\n")

            f.write("\nREVISION TARGETS\n")
            f.write("-" * 104 + "\n")
            targets = cdf[
                cdf["Publication_Action"].isin(["REPLACE", "RELABEL"])
            ]

            if targets.empty:
                f.write("No locked-metric manuscript revisions detected.\n")
            else:
                for _, r in targets.iterrows():
                    f.write(
                        f"{r['Section']} | {r['Location_ID']} | "
                        f"{r['Metric']} -> {r['Outcome']} | "
                        f"rho={r['Found_Rho'] or 'NA'}, "
                        f"p={r['Found_p'] or 'NA'}, "
                        f"q={r['Found_q'] or 'NA'} | "
                        f"{r['Decision']} | action={r['Publication_Action']}\n"
                    )
                    f.write(f"  -> {r['Recommended_Action']}\n")

        f.write("\nSECTION CONSISTENCY MATRIX\n")
        f.write("-" * 104 + "\n")
        for r in matrix:
            f.write(
                f"{r['Metric']:10s} | {r['Section']:11s} | "
                f"N={r['Occurrence_N']} | {r['Status']} | "
                f"rho={r['Found_Rho'] or 'NA'} | "
                f"p={r['Found_p'] or 'NA'} | "
                f"q={r['Found_q'] or 'NA'} | "
                f"actions={r.get('Actions','') or 'NA'}\n"
            )

        f.write("\nV1.1 INTERPRETATION\n")
        f.write("-" * 104 + "\n")
        f.write(
            "1) Only metric-local fragments containing rho/p/q are treated as numerical claims.\n"
            "2) Metric-only prose and figure references are not revision targets.\n"
            "3) Legacy FDR q-values attached to primary claims are explicitly flagged.\n"
            "4) 3D-RMSE and product-sensitivity claims are protected from horizontal-RMSE substitution.\n"
            "5) KEEP/REPLACE/RELABEL/VERIFY_SOURCE are publication actions, not merely parser labels.\n"
            "6) This audit remains read-only; no DOCX content has been changed.\n"
        )

    print(summary.read_text(encoding="utf-8"))


if __name__ == "__main__":
    main()
